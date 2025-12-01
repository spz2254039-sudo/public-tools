# -*- coding: utf-8 -*-
"""
 m o m o _ c h e c k _ v 2
 在 v1 基礎上改為 CSV 匯入、Excel 匯出（保留原 parse_momo_simple 爬蟲邏輯）
"""

import datetime
import html
import re
import time
from typing import Dict, List

import pandas as pd
import requests
from bs4 import BeautifulSoup

try:  # 供舊版 Word 流程使用，預設不在 main() 中呼叫
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
except Exception:  # pragma: no cover - Colab 環境缺少 docx 時允許略過
    Document = None
    Cm = Pt = RGBColor = None


# ➤ 對 MOMO 商品進行 retry + timeout 的簡易爬蟲（保留 v1 邏輯）
def parse_momo_simple(url: str, max_retries: int = 3) -> Dict[str, str]:
    headers = {"User-Agent": "Mozilla/5.0"}
    for attempt in range(1, max_retries + 1):
        try:
            res = requests.get(url, headers=headers, timeout=20)
            res.raise_for_status()
            soup = BeautifulSoup(res.text, "html.parser")

            # 商品名稱
            name_tag = soup.select_one("meta[property='og:title']")
            name = name_tag["content"].strip() if name_tag else "未取得"

            # 品號（多重嘗試）
            prod_no = None
            tag = soup.select_one("#osmPrdNo")
            if tag:
                prod_no = tag.text.strip()
            if not prod_no:
                li_tags = soup.select("li.tvlogo, li.goods-code-container")
                for tag in li_tags:
                    if "品號：" in tag.text:
                        prod_no = tag.text.split("品號：")[-1].strip()
                        break
            if not prod_no:
                meta_code = soup.select_one("meta[name='keywords']")
                if meta_code and "品號：" in meta_code["content"]:
                    prod_no = meta_code["content"].split("品號：")[-1].split(",")[0].strip()
            if not prod_no:
                match = re.search(r"品號[:： ]?\s*(\w+)", soup.get_text())
                prod_no = match.group(1).strip() if match else "未取得"

            return {"商品名稱": name, "品號": prod_no, "網址": url}

        except Exception as e:  # pragma: no cover - 連線異常時輸出錯誤訊息
            print(f"❌ 第 {attempt} 次失敗：{e}")
            if attempt < max_retries:
                time.sleep(3)
            else:
                return {"商品名稱": f"錯誤：{e}", "品號": "錯誤", "網址": url}


# ➤ 擷取網址並處理 &amp; 解碼（保留舊版工具函式）
def extract_urls_from_text(text: str) -> List[str]:
    clean = html.unescape(html.unescape(text))
    urls = re.findall(r"https?://[^\s\]]+", clean)
    return list(dict.fromkeys(urls))


def split_blocks(text: str) -> List[str]:
    parts = re.split(r"(?=收文號：\d+)", text)
    return [part.strip() for part in parts if part.strip()]


def extract_filename(text_block: str) -> str:
    match = re.search(r"收文號：(\d+)", text_block)
    return f"收文號{match.group(1)}-查核報告" if match else "查核報告"


def to_roc_date(date_obj: datetime.date) -> str:
    roc_year = date_obj.year - 1911
    return f"{roc_year}/{date_obj.month}/{date_obj.day}"


def fetch_momo_product(url: str) -> str:
    result = parse_momo_simple(url, max_retries=3)
    name = result.get("商品名稱") if isinstance(result, dict) else None
    if not name or name.startswith("錯誤："):
        return ""
    return name


def load_input_csv() -> pd.DataFrame:
    try:
        from google.colab import files  # type: ignore

        uploaded = files.upload()
        if not uploaded:
            raise FileNotFoundError("No file uploaded.")
        filename = next(iter(uploaded))
    except Exception:
        filename = input("請輸入 CSV 檔名：").strip()
    df = pd.read_csv(filename)
    if "序號" not in df.columns or "商品網址" not in df.columns:
        raise ValueError("CSV 必須包含『序號』與『商品網址』欄位")
    return df


def build_output_rows(df: pd.DataFrame, roc_date: str) -> pd.DataFrame:
    rows: List[Dict[str, str]] = []
    for _, row in df.iterrows():
        url = str(row.get("商品網址", "")).strip()
        product_name = fetch_momo_product(url) if url else ""
        rows.append(
            {
                "編號": row.get("序號", ""),
                "檢查案號": "",
                "查核日期": roc_date,
                "網路名稱/店家名稱": "momo購物網",
                "賣家帳號或拍賣代碼": "",
                "商品名稱": product_name,
                "再查核日期": "",
                "是否下架": "",
                "是否改正": "",
                "調查結果": "",
                "網址/地址": url,
                "商檢標識": "",
                "已宣導": "",
                "已下架": "",
            }
        )
        # ★ 新增這一行：每筆之間停一下，降低被伺服器視為攻擊的機率
        time.sleep(5)
    columns = [
        "編號",
        "檢查案號",
        "查核日期",
        "網路名稱/店家名稱",
        "賣家帳號或拍賣代碼",
        "商品名稱",
        "再查核日期",
        "是否下架",
        "是否改正",
        "調查結果",
        "網址/地址",
        "商檢標識",
        "已宣導",
        "已下架",
    ]
    return pd.DataFrame(rows, columns=columns)


def export_to_excel(df: pd.DataFrame, roc_date: str) -> str:
    filename = f"momo_check_output_ROC{roc_date.replace('/', '')}.xlsx"
    df.to_excel(filename, index=False)
    try:
        from google.colab import files  # type: ignore

        files.download(filename)
    except Exception:
        print(f"已匯出檔案：{filename}")
    return filename


# 舊版流程（raw_text → Word 報告），預設不會在 main() 呼叫
raw_text = ""
# 這裡保留空字串：v1 的範例資料如需使用，請自行填入。


def legacy_word_report():  # pragma: no cover - 僅供需要時手動呼叫
    if Document is None:
        raise ImportError("python-docx 未安裝，無法產生 Word 報告")
    if not raw_text.strip():
        raise ValueError("raw_text 為空，請填入原始內容後再執行")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.27)
        sec.bottom_margin = Cm(1.27)
        sec.left_margin = Cm(1.85)
        sec.right_margin = Cm(1.27)

    blocks = split_blocks(raw_text)
    global_index = 1

    for block in blocks:
        lines = block.strip().split("\n")
        urls = extract_urls_from_text(block)

        results = []
        for url in urls:
            r = parse_momo_simple(url)
            r["編號"] = global_index
            global_index += 1
            results.append(r)
            time.sleep(5)

        for line in lines:
            doc.add_paragraph(line.strip())
            if line.strip().startswith("正本："):
                for row in results:
                    i, name, prod, url = row["編號"], row["商品名稱"], row["品號"], row["網址"]

                    p1 = doc.add_paragraph()
                    r1 = p1.add_run(f"{i}. {name}")
                    r1.font.color.rgb = RGBColor(255, 0, 0)
                    r1.font.size = Pt(12)

                    p2 = doc.add_paragraph()
                    r2 = p2.add_run("(查無商品檢驗標識)")
                    r2.font.color.rgb = RGBColor(255, 0, 0)
                    r2.font.size = Pt(12)

                    doc.add_paragraph(f"品號: {prod}")
                    doc.add_paragraph(f"網址: {url}")
                    doc.add_paragraph()

    fname = extract_filename(blocks[0]) + ".docx"
    doc.save(fname)
    print("✅ 已產出 Word 檔案：", fname)


def main():
    df_in = load_input_csv()
    roc_date = to_roc_date(datetime.date.today())
    df_out = build_output_rows(df_in, roc_date)
    export_to_excel(df_out, roc_date)


if __name__ == "__main__":
    main()
